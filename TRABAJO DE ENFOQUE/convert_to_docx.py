import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Estilo base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Títulos
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('---'):
            # Separador (página nueva o espacio)
            doc.add_page_break()
        elif line.startswith('* '):
            # Listas
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:])
        elif line.startswith('|'):
            # Tablas (Simplificado: omitir o tratar como texto por ahora para evitar errores complejos de parsing)
            # Si es la cabecera de la tabla DAFO, podemos intentar algo básico
            continue 
        else:
            # Párrafo normal
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

    doc.save(docx_path)

def _add_formatted_text(paragraph, text):
    # Manejo básico de negritas **texto**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    md_file = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\Plan_de_Negocio_MurgiSmart_Greens.md'
    docx_file = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\Plan_de_Negocio_MurgiSmart_Greens.docx'
    convert_md_to_docx(md_file, docx_file)
    print(f"Archivo guardado en {docx_file}")
