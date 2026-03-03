import os
from docx import Document
from docx.shared import Pt
import re

def create_final_doc(template_path, md_path, output_path):
    doc = Document(template_path)
    
    # Reemplazar marcadores en la portada
    for para in doc.paragraphs:
        if 'Plantilla trabajo de enfoque' in para.text:
            for run in para.runs:
                run.text = run.text.replace('Plantilla trabajo de enfoque', 'Trabajo de enfoque de Simulación Empresarial')
        if 'Nombre del mdulo profesional' in para.text:
            for run in para.runs:
                run.text = run.text.replace('Nombre del mdulo profesional', 'Simulación Empresarial')
        if 'Ttulo del ndice' in para.text:
            for run in para.runs:
                run.text = run.text.replace('Ttulo del ndice', 'Plan de Negocio: MurgiSmart Greens')

    doc.add_page_break()

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = content.split('---')
    
    for i, sec in enumerate(sections):
        if i == 0 and '# ' in sec: continue # Saltar el titulo H1 inicial
        
        lines = sec.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            
            p = None
            if line.startswith('## '):
                p = doc.add_paragraph(line[3:])
                # Como los estilos pueden fallar, aplicamos formato directo
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = 'Manrope'
            elif line.startswith('### '):
                p = doc.add_paragraph(line[4:])
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(13)
                run.font.name = 'Manrope'
            elif line.startswith('* '):
                p = doc.add_paragraph()
                p.add_run('• ').font.name = 'Manrope'
                _add_formatted_text(p, line[2:])
            elif line.startswith('|'):
                # Ignoramos tablas brutas de MD
                continue
            else:
                p = doc.add_paragraph()
                _add_formatted_text(p, line)
            
            if p:
                for run in p.runs:
                    run.font.name = 'Manrope'
                    if not run.font.size:
                        run.font.size = Pt(11)

    doc.save(output_path)

def _add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Manrope'
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Manrope'
            run.font.size = Pt(11)

if __name__ == "__main__":
    template = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\Plantilla de documentos (1).docx'
    md_source = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\Plan_de_Negocio_MurgiSmart_Greens.md'
    output = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\TRABAJO_FINAL_SIMULACION_MurgiSmart.docx'
    
    create_final_doc(template, md_source, output)
    print(f"Documento final generado en {output}")
