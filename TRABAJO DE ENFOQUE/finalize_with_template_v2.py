import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def create_final_doc_v2(template_path, md_path, output_path):
    doc = Document(template_path)
    
    # Definimos el AZUL corporativo detectado (003594)
    CORP_BLUE = RGBColor(0x00, 0x35, 0x94)
    # Color secundario/naranja segun P5 (FF5B41)
    CORP_ORANGE = RGBColor(0xFF, 0x5B, 0x41)

    # 1. Ajustar Portada
    for para in doc.paragraphs:
        if 'Plantilla trabajo de enfoque' in para.text:
            para.text = ''
            run = para.add_run('TRABAJO DE ENFOQUE DE SIMULACIÓN EMPRESARIAL')
            run.font.name = 'Manrope'
            run.font.size = Pt(26)
            run.font.bold = True
            run.font.color.rgb = CORP_BLUE
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if 'Nombre del mdulo profesional' in para.text:
            para.text = ''
            run = para.add_run('SIMULACIÓN EMPRESARIAL')
            run.font.name = 'Manrope'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = CORP_ORANGE
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        if 'Ttulo del ndice' in para.text:
            para.text = ''
            run = para.add_run('ÍNDICE DE CONTENIDOS')
            run.font.name = 'Manrope'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = CORP_ORANGE

    # 2. Insertar Índice (TOC)
    # Nota: Word pide actualizar el campo al abrir, pero lo dejamos listo.
    _add_table_of_contents(doc)
    doc.add_page_break()

    # 3. Leer y Procesar Contenido
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = content.split('---')
    
    for i, sec in enumerate(sections):
        if i == 0 and '# ' in sec: continue # Saltar H1 duplicado
        
        lines = sec.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            
            p = None
            if line.startswith('## '):
                p = doc.add_paragraph(line[3:])
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(18) # Heading 1 size
                run.font.name = 'Manrope'
                run.font.color.rgb = CORP_BLUE
            elif line.startswith('### '):
                p = doc.add_paragraph(line[4:])
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(14) # Heading 2 size
                run.font.name = 'Manrope'
                run.font.color.rgb = CORP_BLUE
            elif line.startswith('* '):
                p = doc.add_paragraph()
                p.add_run('• ').font.name = 'Manrope'
                _add_formatted_text(p, line[2:], CORP_BLUE)
            else:
                p = doc.add_paragraph()
                _add_formatted_text(p, line, CORP_BLUE)
            
            if p:
                for run in p.runs:
                    run.font.name = 'Manrope'
                    if not run.font.size:
                        run.font.size = Pt(11)

    doc.save(output_path)

def _add_formatted_text(paragraph, text, blue_color):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Manrope'
        run.font.size = Pt(11)

def _add_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

if __name__ == "__main__":
    template = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\Plantilla de documentos (1).docx'
    md_source = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\Plan_de_Negocio_MurgiSmart_Greens.md'
    output = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\ENTREGA_DEFINITIVA_SIMULACION_MurgiSmart.docx'
    
    create_final_doc_v2(template, md_source, output)
    print(f"Documento final (V2) generado en {output}")
