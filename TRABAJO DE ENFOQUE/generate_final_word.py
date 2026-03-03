import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def create_ultimate_doc_v3(template_path, md_path, output_path):
    doc = Document(template_path)
    
    CORP_BLUE = RGBColor(0x00, 0x35, 0x94)
    CORP_ORANGE = RGBColor(0xFF, 0x5B, 0x41)

    # 1. Preparar Portada
    for para in doc.paragraphs:
        if 'Plantilla trabajo de enfoque' in para.text:
            para.text = ''
            run = para.add_run('PROYECTO: PLAN DE NEGOCIO')
            run.font.name = 'Manrope'
            run.font.size = Pt(26)
            run.font.bold = True
            run.font.color.rgb = CORP_BLUE
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if 'Nombre del mdulo profesional' in para.text:
            para.text = ''
            run = para.add_run('SIMULACIÓN EMPRESARIAL - MURGI SL')
            run.font.name = 'Manrope'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = CORP_ORANGE
            
        if 'Ttulo del ndice' in para.text:
            para.text = ''
            # Este marcador lo quitamos para poner el índice real debajo

    # 2. Recolectar Headings para el ÍNDICE estático
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    headings = []
    for line in lines:
        if line.startswith('## '):
            headings.append(line[3:].strip())
        elif line.startswith('### '):
            headings.append("  " + line[4:].strip())

    # 3. Crear el ÍNDICE estático y visible
    index_para = doc.add_paragraph()
    run = index_para.add_run('ÍNDICE DE CONTENIDOS')
    run.font.name = 'Manrope'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = CORP_BLUE
    index_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for h in headings:
        p = doc.add_paragraph()
        run = p.add_run(h)
        run.font.name = 'Manrope'
        run.font.size = Pt(11)
        # Añadir puntos suspensivos si queremos estetica de indice
        p.add_run(' ........................................................................')

    doc.add_page_break()

    # 4. Procesar el Contenido Real
    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            if '---' not in line:
                cells = [cell.strip() for cell in line.split('|')]
                if len(cells) > 1:
                    table_data.append(cells[1:-1])
            continue
        else:
            if in_table:
                if table_data:
                    _add_word_table_manual(doc, table_data, CORP_BLUE)
                in_table = False
                table_data = []

        if not line: continue
        if line.startswith('# '): continue
        
        if line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = 'Manrope'
            run.font.color.rgb = CORP_BLUE
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Manrope'
            run.font.color.rgb = CORP_BLUE
        elif line.startswith('* '):
            p = doc.add_paragraph()
            p.add_run('• ').font.name = 'Manrope'
            _add_formatted_text(p, line[2:])
        elif line.startswith('---'):
            doc.add_page_break()
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

    doc.save(output_path)

def _add_formatted_text(paragraph, text):
    # Regex para capturar tanto **texto** como *texto* y aplicarle negrita
    # El usuario pidió: "las palabras en negrita, cuando hay asteriscos"
    pattern = r'(\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, text)
    for part in parts:
        if (part.startswith('**') and part.endswith('**')) or (part.startswith('*') and part.endswith('*')):
            # Quitamos los asteriscos de los bordes
            clean_text = part.replace('**', '').replace('*', '')
            run = paragraph.add_run(clean_text)
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Manrope'
        run.font.size = Pt(11)

def _add_word_table_manual(doc, data, header_color):
    if not data: return
    rows_count = len(data)
    cols_count = len(data[0])
    table = doc.add_table(rows=rows_count, cols=cols_count)
    
    # Bordes manuales
    tbl = table._tbl
    tblPr = tbl.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    for r_idx, row_data in enumerate(data):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < cols_count:
                cell = table.cell(r_idx, c_idx)
                para = cell.paragraphs[0]
                _add_formatted_text(para, cell_text)
                if r_idx == 0:
                    for run in para.runs:
                        run.font.color.rgb = header_color
                        run.bold = True

if __name__ == "__main__":
    template = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\Plantilla de documentos (1).docx'
    md_source = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\Plan_de_Negocio_Murgi_SL.md'
    output = r'f:\17. SIMULACION EMPRESARIAL\TRABAJO DE ENFOQUE\TRABAJO_FINAL_MURGI_SL_V3.docx'
    
    create_ultimate_doc_v3(template, md_source, output)
    print(f"Documento definitivo generado en {output}")
