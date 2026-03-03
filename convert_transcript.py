import json
import os
import argparse
from docx import Document

def convert_transcript_to_docx(input_path, output_path):
    # Check if input file exists
    if not os.path.exists(input_path):
        print(f"Error: Input file not found at {input_path}")
        return

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {e}")
        return
    except Exception as e:
        print(f"Error reading file: {e}")
        return

    # Extract text segments
    full_text = ""
    # The structure seems to be: data['events'] -> list of events -> each event has 'segs' -> list of segments -> 'utf8' field
    
    events = data.get('events', [])
    if not events:
        print("No events found in JSON.")
        return

    print(f"Found {len(events)} events.")

    for event in events:
        segs = event.get('segs', [])
        if segs:
            for seg in segs:
                utf8_text = seg.get('utf8', '')
                if utf8_text:
                    full_text += utf8_text

    # Create Docx
    doc = Document()
    doc.add_heading(f'Transcripción - {os.path.basename(input_path)}', 0)
    
    paragraphs = full_text.split('\n')
    
    for para in paragraphs:
        if para.strip(): # Only add non-empty paragraphs
            doc.add_paragraph(para.strip())

    try:
        doc.save(output_path)
        print(f"Successfully created: {output_path}")
    except Exception as e:
        print(f"Error saving DOCX: {e}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert JSON transcript to DOCX")
    parser.add_argument("input_file", help="Path to the input JSON/TXT file")
    parser.add_argument("output_file", help="Path to the output DOCX file")
    
    args = parser.parse_args()
    
    convert_transcript_to_docx(args.input_file, args.output_file)
